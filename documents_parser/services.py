from datetime import datetime
from .models import ExtractedData  # Import your Django model
from .excel_controller import *
from .pdf_controller import *
from .github_controller import *
import uuid

from docx import Document
from io import BytesIO
from django.utils.html import escape
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX  

def process_uploaded_file(file_path, uploaded_file_name, tag_names, user):
    results_by_tag_exact = {tag: [] for tag in tag_names}
    results_by_tag_partial = {tag: [] for tag in tag_names}
    file_type=None
    if uploaded_file_name.endswith('.pdf'):
        file_type="PDF"
        # Process PDF files
        text_data = read_pdf_file(file_path)
        for tag in tag_names:
            result_exact, result_partial = search_pdf(text_data, tag, uploaded_file_name, "PDF", user)
            results_by_tag_exact[tag] += result_exact
            results_by_tag_partial[tag] += result_partial

    elif uploaded_file_name.endswith('.docx'):
        file_type="Word"
        text_data = read_doc_file(file_path)
        for tag in tag_names:
            result_exact, result_partial = search_pdf(text_data, tag, uploaded_file_name, "Word", user)
            results_by_tag_exact[tag] += result_exact
            results_by_tag_partial[tag] += result_partial

    elif uploaded_file_name.endswith('.xls') or uploaded_file_name.endswith('.xlsx'):
        file_type="Excel"
        for tag in tag_names:
            result_exact, result_partial = extract_text_from_excel(file_path, tag, uploaded_file_name, user)
            results_by_tag_exact[tag] += result_exact
            results_by_tag_partial[tag] += result_partial

    elif uploaded_file_name.endswith('.zip'):
        file_type="Github"
        zip_file_name = uploaded_file_name.split(".")[0]
        data_dict = read_all_files_in_directory(file_path)
        for tag in tag_names:
            result_exact, result_partial = search_github(tag, zip_file_name, data_dict)
            results_by_tag_exact[tag] += result_exact
            results_by_tag_partial[tag] += result_partial

    else:
        raise ValueError(f"Unsupported file format: {uploaded_file_name}")

    return file_type, results_by_tag_exact, results_by_tag_partial



def save_results_to_db(results,results_partial, file_name, file_type, user):

    search_id = uuid.uuid4()
    print("UUID",search_id)

    saved_results_exact = [] 
    saved_results_partial = []

    for tag, result_list in results.items():
        for result in result_list:
            saved_result = ExtractedData.objects.create(
                search_id=search_id,  # Assign the generated search ID
                source_file_name=result["Source File Name"],
                file_type=result["File Type"],
                tag_searched=tag,
                block_record=result['Block/Record'],
                location_of_tag=result['Location of the Tag'],
                # date_of_search is automatically set with auto_now_add=True
                search_author=user,
                match_type="exact",
                other=result.get('Other', ''),
                # line_id will be automatically generated in the save method if not provided
            )
            # Append the saved instance to the list
            saved_results_exact.append(saved_result)
    
    
    for tag, result_list in results_partial.items():
        for result in result_list:
            saved_result = ExtractedData.objects.create(
                search_id=search_id,  # Assign the generated search ID
                source_file_name=result["Source File Name"],
                file_type=result["File Type"],
                tag_searched=tag,
                block_record=result['Block/Record'],
                location_of_tag=result['Location of the Tag'],
                # date_of_search is automatically set with auto_now_add=True
                search_author=user,
                match_type="partial",
                other=result.get('Other', ''),
                # line_id will be automatically generated in the save method if not provided
            )
            # Append the saved instance to the list
            saved_results_partial.append(saved_result)


    return saved_results_exact,saved_results_partial



def find_records(source_file_name=None, file_type=None, tag_searched=None, start_date=None, end_date=None, search_author=None):

    records = ExtractedData.objects.all()

    # Apply filters based on provided arguments
    if source_file_name:
        records = records.filter(source_file_name__icontains=source_file_name)

    if file_type:
        records = records.filter(file_type__iexact=file_type)

    if tag_searched:
        records = records.filter(tag_searched__icontains=tag_searched)

    if start_date:
        start_date_obj = datetime.strptime(start_date, "%Y-%m-%d").date()
        records = records.filter(date_of_search__gte=start_date_obj)

    if end_date:
        end_date_obj = datetime.strptime(end_date, "%Y-%m-%d").date()
        records = records.filter(date_of_search__lte=end_date_obj)

    if search_author:
        records = records.filter(search_author__icontains=search_author)

    return records





def highlight_text(paragraph, tag):
    """
    Highlights all occurrences of a tag in a paragraph with yellow color.
    """
    # Split paragraph text around the tag
    parts = paragraph.text.split(tag)

    if len(parts) <= 1:
        return  # Tag not found, nothing to highlight

    # Clear existing paragraph text
    paragraph.clear()

    # Reconstruct paragraph with highlighted tag
    for i, part in enumerate(parts):
        if i > 0:
            # Create a run for the highlighted tag
            run = paragraph.add_run(tag)
            run.font.highlight_color = 'yellow'  # Apply text highlight color

        # Add the normal text
        paragraph.add_run(part)



def highlight_text(paragraph, tag):
    """
    Highlights all occurrences of a tag in a paragraph with yellow color.
    """
    # Split paragraph text around the tag
    parts = paragraph.text.split(tag)

    if len(parts) <= 1:
        return  # Tag not found, nothing to highlight

    # Clear existing paragraph text
    paragraph.clear()

    # Reconstruct paragraph with highlighted tag
    for i, part in enumerate(parts):
        if i > 0:
            # Create a run for the highlighted tag
            run = paragraph.add_run(tag)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Apply text highlight color using WD_COLOR_INDEX

        # Add the normal text
        paragraph.add_run(part)


def export_search_results_to_word(search_results, file_name):
    """
    Export search results to a Word document, creating separate tables for exact and partial matches.
    """
    document = Document()

    # Formatting the title with document name and current date
    current_date = datetime.now().strftime("%B %d, %Y")
    document_title = f'Search result for {file_name} on {current_date}'
    document.add_heading(document_title, level=1)

    headers = ['Source File Name', 'File Type', 'Tag Searched', 'Block/Record', 'Location of the Tag', 'Date and Time', 'Author Name', 'Other']

    # Create a section for Exact Matches
    document.add_heading('Exact Matches', level=2)
    table_exact = document.add_table(rows=1, cols=len(headers))
    table_exact.style = 'Table Grid'
    hdr_cells_exact = table_exact.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells_exact[i].text = header

    # Fill the Exact Matches table
    for result in search_results.filter(match_type='exact'):
        row_cells = table_exact.add_row().cells
        row_cells[0].text = result.source_file_name
        row_cells[1].text = result.file_type
        row_cells[2].text = result.tag_searched
        row_cells[3].text = result.block_record
        row_cells[4].text = result.location_of_tag
        row_cells[5].text = result.date_of_search.strftime("%Y-%m-%d")
        row_cells[6].text = result.search_author
        row_cells[7].text = result.other

        # Highlight the tag in the 'Block/Record' cell
        highlight_text(row_cells[3].paragraphs[0], result.tag_searched)

    # Create a section for Partial Matches
    document.add_heading('Partial Matches', level=2)
    table_partial = document.add_table(rows=1, cols=len(headers))
    table_partial.style = 'Table Grid'
    hdr_cells_partial = table_partial.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells_partial[i].text = header

    # Fill the Partial Matches table
    for result in search_results.filter(match_type='partial'):
        row_cells = table_partial.add_row().cells
        row_cells[0].text = result.source_file_name
        row_cells[1].text = result.file_type
        row_cells[2].text = result.tag_searched
        row_cells[3].text = result.block_record
        row_cells[4].text = result.location_of_tag
        row_cells[5].text = result.date_of_search.strftime("%Y-%m-%d")
        row_cells[6].text = result.search_author
        row_cells[7].text = result.other

        # Highlight the tag in the 'Block/Record' cell
        highlight_text(row_cells[3].paragraphs[0], result.tag_searched)

    # Save document to a BytesIO object
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream


def append_dicts(dict1, dict2):
    # Ensure both dictionaries have the same keys
    if dict1.keys() != dict2.keys():
        raise ValueError("Both dictionaries must have the same keys.")

    # Create a new dictionary to store the appended values
    result_dict = {}

    for key in dict1.keys():
        # Check if the value is a list; if not, convert to list for appending
        value1 = dict1[key] if isinstance(dict1[key], list) else [dict1[key]]
        value2 = dict2[key] if isinstance(dict2[key], list) else [dict2[key]]

        # Append values from dict2 to dict1 under the same key
        result_dict[key] = value1 + value2
    
    

    return result_dict